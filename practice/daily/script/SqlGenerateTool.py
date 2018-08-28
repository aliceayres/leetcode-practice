import xlrd
from docx import Document
import os

class Table:
    def __init__(self,name,comment,fields,primaryKey):
        '''
        :param name: 表名
        :param comment: 表注释
        :param fields: 字段 List
        :param primaryKey: 主键名
        '''
        self.name = name
        self.comment = comment
        self.fields = fields
        self.primaryKey = primaryKey

    def createSql(self):
        '''
        获取建表SQL语句
        :return: 建表sql
        '''
        sql = '# 建表[ ' + self.comment +' - '+self.name + ' ]的SQL语句\n'
        sql += 'DROP TABLE IF EXISTS ' + self.name + ';\n'
        sql += 'CREATE TABLE ' + self.name + ' (\n'
        for i in range(0,len(self.fields)):
            field = self.fields[i]
            sql += field.name + ' ' + field.type
            if field.notnull:
                sql += ' NOT'
            sql += ' NULL'
            if field.comment != '':
                sql += ' COMMENT \'' + field.comment + '\''
            if i != len(self.fields) - 1:
                sql += ',\n'
        sql += '\n)'
        if self.comment != '':
            sql += 'comment = \''+self.comment+'\''
        sql += ';\n'
        if self.primaryKey != '':
            sql += 'ALTER TABLE ' + self.name + ' ADD PRIMARY KEY(' + self.primaryKey + ');\n'
        return sql

class Field:
    def __init__(self,name,comment,type,notnull):
        '''
        :param name: 字段名称
        :param comment: 字段注释
        :param type: 字段类型
        :param notnull: 字段是否非空
        '''
        self.name = name
        self.comment = comment
        self.type = type
        self.notnull = notnull
        self.preProcessType()

    def preProcessType(self):
        '''
        预处理字段类型 Oracle 2 MySQL
        :return:
        '''
        self.type = self.type.upper()
        self.type = self.type.replace(' ', '')
        if self.type.find('CLOB') > -1:
            self.type = self.type.replace('CLOB', 'TEXT')
        if self.type.find('NVARCHAR2') > -1:
            self.type = self.type.replace('NVARCHAR2', 'VARCHAR')
        if self.type.find('VARCHAR2') > -1:
            self.type = self.type.replace('VARCHAR2', 'VARCHAR')
        if self.type.find('NUMBER') > -1:
            self.type = self.type.replace('NUMBER', 'NUMERIC')
        if self.type.find('NUMERIC(138)') > -1:
            self.type = self.type.replace('NUMERIC(138)', 'NUMERIC(65)')
        if self.type.find('DATE') > -1:
            self.type = self.type.replace('DATE', 'DATETIME')
        if self.type.find('，') > -1:
            self.type = self.type.replace('，', ',')

class TableTool:
    def tablesToSqlFile(self,root,tables):
        '''
        存储为SQL文件
        :param root: Sql文件存放根目录
        :param tables: 表 List
        :return:
        查询数据库中表数量：SELECT COUNT( * ) FROM information_schema.tables WHERE TABLE_SCHEMA = 'dbo'
        默认汇总文件名称：all_table.sql
        '''
        if not os.path.exists(root):
            os.makedirs(root)
        total_filename = 'all_table.sql'
        total = ''
        for tb in tables:
            sql = tb.createSql()
            total += sql +'\n'
            file = open(root + tb.name.lower() + '.sql', 'w', encoding="utf-8")
            file.write(sql)
            file.close()
        total_file = open(root + total_filename, 'w', encoding="utf-8")
        total_file.write(total)
        total_file.close()

    def tablesFromExcel(self,filename):
        '''
        Excel获取表结构
        :param filename: excel 文件名
        :return: Table list
        excel文件格式：首行title —— 序号|字段备注|字段名称|字段类型|非空‘Y’|备注
        '''
        workbook = xlrd.open_workbook(filename)
        all_sheets = workbook.sheet_names()
        tables = []
        for i in range(0,len(all_sheets)):
            table_name = all_sheets[i]
            table_comment = ''
            sheet = workbook.sheet_by_name(all_sheets[i]) # workbook.sheet_by_index(1)
            rows = sheet.nrows
            fields = []
            for j in range(1, rows):
                field_name = sheet.row_values(j)[2]
                field_type = sheet.row_values(j)[3]
                field_comment = sheet.row_values(j)[1]
                not_null = sheet.row_values(j)[4] == 'Y'
                field = Field(field_name,field_comment,field_type,not_null)
                fields.append(field)
            table = Table(table_name,table_comment,fields,fields[0].name)
            tables.append(table)
        return tables

    def tablesFromDocx(self,filename):
        '''
        Word文档获取表结构
        :param filename:
        :return:
        '''
        document = Document(filename)
        all_tables = document.tables
        tables = []
        alltypes = {}
        for i in range(1, len(all_tables)):
            table = all_tables[i]
            style = 0
            # 判断表格类型
            if len(table.rows) >= 5:
                if table.rows[4].cells[0].text == '序号':
                    style = 1
            # 表格类型 0 处理
            if style != 1:
                # continue
                title = []
                for tcell in table.rows[0].cells:
                    if len(title) == 0 or (len(title) > 0 and tcell.text != title[-1]):
                        title.append(tcell.text)
                table_comments = title[0]
                table_name = title[1]
                print('表注释-[%s],表名-[%s]' % (table_comments, table_name))
                fields = []
                for k in range(2, len(table.rows)):
                    row = table.rows[k]
                    cells = []
                    for cell in row.cells:
                        if len(cells) == 0 or (len(cells) > 0 and cell.text != cells[-1]):
                            cells.append(cell.text)
                    if len(cells) == 1:
                        continue
                    print(cells)
                    name = cells[0].upper()
                    comment = cells[3]
                    type = cells[1].upper()
                    notnull = False
                    alltypes.setdefault(type)
                    field = Field(name, comment, type, notnull)
                    fields.append(field)
                tb = Table(table_name, table_comments, fields, '')
                tables.append(tb)
            # 表格类型 1 处理
            if style == 1:
                # continue
                # 获取表名、表备注
                title = []
                for tcell in table.rows[0].cells:
                    if len(title) == 0 or (len(title) > 0 and tcell.text != title[-1]):
                        title.append(tcell.text)
                table_comments = title[1]
                table_name = title[3]
                print('表注释-[%s],表名-[%s]' % (table_comments, table_name))
                fields = []
                for k in range(0, len(table.rows)):
                    row = table.rows[k]
                    if k < 5:
                        continue
                    if row.cells[0].text == '主键' or row.cells[0].text == '索引':
                        break
                    cells = []
                    for cell in row.cells:
                        if len(cells) == 0 or (len(cells) > 0 and cell.text != cells[-1]):
                            if len(cells) < 5:
                                cells.append(cell.text)
                            repeated = {'ID','DNS1','DNS2','CPU'}
                            if table_name != 'IDC_RHS_ORDER_UPDATE_HIS':
                                if cell.text in repeated and len(cells) == 2:
                                    cells.append(cell.text)
                    if len(cells) == 1:
                        continue
                    print(cells)
                    name = cells[2].upper()
                    comment = cells[1]
                    type = cells[3].upper()
                    notnull = cells[4] == 'Y'
                    alltypes.setdefault(type)
                    field = Field(name,comment,type,notnull)
                    fields.append(field)
                tb = Table(table_name,table_comments,fields,fields[0].name)
                tables.append(tb)
        # for tps in alltypes.keys():
        #     print(tps)
        return tables

class Solution:
    def sqlFromOfficeFile(self,root,filename):
        tool = TableTool()
        if filename.find('.xlsx') > -1:
            tables = tool.tablesFromExcel(filename)
        if filename.find('.docx') > -1:
            tables = tool.tablesFromDocx(filename)
        tool.tablesToSqlFile(root, tables)

    def process(self):
        root = 'D:\Ayres\idcm_sql\\'
        filename = 'D:\Ayres\dbo.docx'
        self.sqlFromOfficeFile(root,filename)

slt = Solution()
slt.process()